"""Generate YFM2 Tower Setup Guide as .docx for Google Drive / Word."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "YFM2-Tower-Setup-Guide.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x21, 0x24)


def add_callout(doc: Document, title: str, body: str, kind: str = "info") -> None:
    colors = {
        "info": (RGBColor(0x17, 0x4E, 0xA6), "E8F0FE"),
        "warn": (RGBColor(0x7C, 0x4A, 0x00), "FEF7E0"),
        "tip": (RGBColor(0x13, 0x73, 0x33), "E6F4EA"),
    }
    title_color, _ = colors.get(kind, colors["info"])
    p = doc.add_paragraph()
    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.color.rgb = title_color
    r2 = p.add_run(body)
    r2.font.size = Pt(10)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("YFM2 Admin · Operations Guide")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Getting the Tower Live")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x0D, 0x0F, 0x14)

    sub = doc.add_paragraph("Step-by-step setup for the Frontier eligibility pipeline")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

    meta = doc.add_paragraph("Live site: yfm2-theta.vercel.app")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x80, 0x86, 0x8B)

    doc.add_paragraph()

    # TOC
    doc.add_heading("Contents", level=2)
    for item in [
        "Overview — what runs where",
        "One-time setup on the tower PC",
        "Starting the tower (every session)",
        "Daily checklist",
        "Running a test job",
        "Troubleshooting",
        "Quick reference commands",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # Section 1
    doc.add_heading("1. Overview — what runs where", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    headers = ["Machine", "Role"]
    rows = [
        ("Dev PC", "Edit website code in Cursor. Optional local testing with npm run dev."),
        ("Tower PC", "Runs the API server, Frontier bots, and Chrome. Must stay on during jobs."),
        ("Vercel", "Hosts the website only. Never runs zip checker or qualifier bots."),
        ("Tailscale Funnel", "Public HTTPS URL so the live website can reach your tower PC."),
    ]
    for i, htext in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = htext
        set_cell_shading(cell, "E8F0FE")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (a, b) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    add_callout(
        doc,
        "Key idea",
        "The website sends job orders to the tower. The tower runs the bots and sends progress back. "
        "If either tower window is closed, the site shows Tower offline.",
        "info",
    )

    # Section 2
    doc.add_heading("2. One-time setup on the tower PC", level=1)
    doc.add_paragraph("Do this once, or again after updating tower-server or bot files on your dev PC.")

    doc.add_heading("Step 2.1 — Confirm Google Drive folder layout", level=2)
    doc.add_paragraph("Files sync via Google Drive to G:\\My Drive\\")
    add_code_block(
        doc,
        "G:\\My Drive\\\n"
        "├── bot\\\n"
        "│   ├── frontier_zipcheck_v2.py      ← Zip checker (phase 1)\n"
        "│   └── frontier_checker60.py       ← Qualifier (phase 2)\n"
        "└── tower-server\\\n"
        "    ├── main.py\n"
        "    ├── start.bat\n"
        "    ├── requirements.txt\n"
        "    └── .env                          ← you create this",
    )

    doc.add_heading("Step 2.2 — Install tower-server dependencies", level=2)
    doc.add_paragraph("Open Command Prompt. Run one command at a time:")
    add_code_block(doc, 'cd /d "G:\\My Drive\\tower-server"\npip install -r requirements.txt')

    doc.add_heading("Step 2.3 — Create the .env file", level=2)
    add_code_block(doc, 'cd /d "G:\\My Drive\\tower-server"\ncopy .env.example .env')
    add_code_block(
        doc,
        "TOWER_BOT_DIR=G:\\My Drive\\bot\n"
        "TOWER_PYTHON=python\n"
        "TOWER_JOBS_DIR=./jobs\n"
        "TOWER_PORT=8787",
    )
    add_callout(
        doc,
        "Important",
        "TOWER_PYTHON=python must be the Python with bot packages (pandas, selenium, "
        "undetected-chromedriver). Use system Python — not a tiny venv with only FastAPI.",
        "warn",
    )

    doc.add_heading("Step 2.4 — Install Tailscale", level=2)
    doc.add_paragraph("Download from https://tailscale.com, install, sign in, leave it running.", style="List Bullet")

    doc.add_heading("Step 2.5 — Set Vercel environment variable", level=2)
    doc.add_paragraph("Variable name: VITE_TOWER_API_URL")
    doc.add_paragraph("Example value: https://desktop-xxxxx.tail5a5189.ts.net")
    doc.add_paragraph("No trailing slash. Do not add /health at the end.")
    doc.add_paragraph("Vercel → yfm2 project → Settings → Environment Variables → redeploy after changes.")

    doc.add_page_break()

    # Section 3
    doc.add_heading("3. Starting the tower (every session)", level=1)
    add_callout(
        doc,
        "You need TWO windows open",
        "Closing either window takes the tower offline. Keep both open while jobs should run.",
        "tip",
    )

    doc.add_heading("Window 1 — Start the API server", level=2)
    add_code_block(doc, 'cd /d "G:\\My Drive\\tower-server"\npython main.py')
    doc.add_paragraph('Expected: "Uvicorn running on http://0.0.0.0:8787"')
    doc.add_paragraph("Leave this window open.")
    doc.add_paragraph("Test in browser on tower PC:")
    add_code_block(doc, "http://localhost:8787/health")

    doc.add_heading("Window 2 — Expose tower to the internet", level=2)
    add_code_block(doc, "tailscale funnel 8787")
    doc.add_paragraph("Copy the HTTPS URL Tailscale prints. Leave this window open.")
    add_callout(
        doc,
        "Normal behavior",
        'Visiting the bare funnel URL may show {"detail":"Not Found"}. Use /health or /api/ paths instead.',
        "info",
    )

    doc.add_heading("Step 3.3 — Update Vercel with funnel URL", level=2)
    doc.add_paragraph("Set VITE_TOWER_API_URL to your funnel URL and redeploy if changed.")

    doc.add_heading("Step 3.4 — Confirm on live website", level=2)
    for step in [
        "Open https://yfm2-theta.vercel.app",
        "Log in as admin",
        "Go to Eligibility tab",
        "Top right should show Tower online (green)",
    ]:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph(
        "If Tower offline (red): click badge → dropdown → Check now after fixing both windows."
    )

    # Section 4
    doc.add_heading("4. Daily checklist", level=1)
    for item in [
        "Tower PC is on and logged in",
        "Google Drive finished syncing",
        "Window 1: python main.py running",
        "Window 2: tailscale funnel 8787 running",
        "Live site shows Tower online",
        "Optional: localhost:8787/health returns ok",
    ]:
        p = doc.add_paragraph(f"☐  {item}")

    # Section 5
    doc.add_heading("5. Running a test job", level=1)
    for step in [
        "Eligibility → Frontier",
        "Enter 5-digit ZIP",
        "Click Run Pipeline",
        "Watch Zip Checker then Qualifier",
        "Use Tower output panel for live logs",
    ]:
        doc.add_paragraph(step, style="List Number")
    add_callout(
        doc,
        "Retry tip",
        "If zip checker succeeded but qualifier failed, use Retry qualifier — it skips zip check.",
        "tip",
    )

    doc.add_page_break()

    # Section 6
    doc.add_heading("6. Troubleshooting", level=1)
    t2 = doc.add_table(rows=7, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Problem"
    t2.rows[0].cells[1].text = "Solution"
    set_cell_shading(t2.rows[0].cells[0], "E8F0FE")
    set_cell_shading(t2.rows[0].cells[1], "E8F0FE")
    fixes = [
        ("Tower offline on site", "Restart both windows. Match Vercel URL to current funnel URL."),
        ("cd fails", 'Use cd /d "G:\\My Drive\\tower-server" with /d and quotes.'),
        ("PowerShell >> prompt", "Unclosed quote — Ctrl+C, new window, one command at a time."),
        ("Qualifier fails at browser", "Fix Chrome / version_main in frontier_checker60.py."),
        ("URL not configured", "Set VITE_TOWER_API_URL on Vercel and redeploy."),
        ("Funnel URL changed", "Re-run tailscale funnel 8787 and update Vercel."),
    ]
    for i, (a, b) in enumerate(fixes, start=1):
        t2.rows[i].cells[0].text = a
        t2.rows[i].cells[1].text = b

    # Section 7
    doc.add_heading("7. Quick reference commands", level=1)
    doc.add_paragraph("Window 1 — API server", style="List Bullet")
    add_code_block(doc, 'cd /d "G:\\My Drive\\tower-server"\npython main.py')
    doc.add_paragraph("Window 2 — Public URL", style="List Bullet")
    add_code_block(doc, "tailscale funnel 8787")
    doc.add_paragraph("Local health check", style="List Bullet")
    add_code_block(doc, "http://localhost:8787/health")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("YFM2 Tower Setup Guide · frontier_checker60")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x86, 0x8B)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
